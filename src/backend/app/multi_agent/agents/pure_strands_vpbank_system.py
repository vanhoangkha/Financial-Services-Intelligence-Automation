"""
VPBank K-MULT Agent Studio - Pure Strands Agents Implementation
Clean architecture using existing VPBank services and nodes
"""

from strands import Agent, tool
from strands.models import BedrockModel
import boto3
import asyncio
import json
import logging
import os
import re
import ssl
import urllib3
import time
from typing import Dict, Any, Optional, List
from datetime import datetime
from uuid import uuid4

# Import VPBank configurations
from app.multi_agent.config import (
    AWS_ACCESS_KEY_ID,
    AWS_SECRET_ACCESS_KEY,
    AWS_BEDROCK_REGION,
    DEFAULT_MODEL_NAME,
    MODEL_MAPPING,
    VERIFY_HTTPS
)

# Import existing VPBank services
from app.multi_agent.services.text_service import TextSummaryService
from app.multi_agent.services.compliance_service import ComplianceValidationService
from app.multi_agent.helpers.improved_pdf_extractor import ImprovedPDFExtractor

logger = logging.getLogger(__name__)

# ================================
# SSL CONFIGURATION FOR STRANDS
# ================================
if not VERIFY_HTTPS:
    # Disable SSL warnings
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    # Create unverified SSL context globally
    ssl._create_default_https_context = ssl._create_unverified_context
    logger.info("SSL verification disabled for Strands Agent")

# ================================
# AWS BEDROCK MODEL CONFIGURATION
# ================================

BEDROCK_MODEL_ID = MODEL_MAPPING.get(DEFAULT_MODEL_NAME, "us.anthropic.claude-3-7-sonnet-20250219-v1:0")

if DEFAULT_MODEL_NAME == "claude-37-sonnet":
    BEDROCK_MODEL_ID = "us.anthropic.claude-3-7-sonnet-20250219-v1:0"

logger.info(f"[PURE_STRANDS] Using model: {BEDROCK_MODEL_ID}")

# Create BedrockModel for Strands - Simple configuration
boto_session = boto3.Session(
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
    region_name=AWS_BEDROCK_REGION,
)

bedrock_model = BedrockModel(
    model_id=BEDROCK_MODEL_ID,
    boto_session=boto_session,
    temperature=0.7,
    top_p=0.9,
    streaming=True
)

# ================================
# ASYNC HELPER FUNCTION
# ================================

def _run_async_safely(async_func):
    """
    Safely run async function in sync context
    """
    try:
        logger.info(f"[ASYNC_WRAPPER] Starting async function: {async_func.__name__ if hasattr(async_func, '__name__') else 'unknown'}")
        
        # Try to get current event loop
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # We're in an async context, use thread executor
            import concurrent.futures
            import threading
            
            logger.info("[ASYNC_WRAPPER] Running in thread executor due to existing event loop")
            
            def run_in_thread():
                # Create new event loop for this thread
                new_loop = asyncio.new_event_loop()
                asyncio.set_event_loop(new_loop)
                try:
                    logger.info("[ASYNC_WRAPPER] Created new event loop in thread")
                    result = new_loop.run_until_complete(async_func())
                    logger.info("[ASYNC_WRAPPER] Async function completed successfully in thread")
                    return result
                except Exception as e:
                    logger.error(f"[ASYNC_WRAPPER] Error in thread execution: {e}")
                    raise e
                finally:
                    new_loop.close()
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(run_in_thread)
                try:
                    result = future.result()  # No timeout - same as original service
                    logger.info("[ASYNC_WRAPPER] Thread executor completed successfully")
                    return result
                except Exception as e:
                    logger.error(f"[ASYNC_WRAPPER] Thread executor failed: {e}")
                    raise e
        else:
            # No running loop, safe to use asyncio.run
            logger.info("[ASYNC_WRAPPER] Running with asyncio.run (no existing loop)")
            result = asyncio.run(async_func())
            logger.info("[ASYNC_WRAPPER] asyncio.run completed successfully")
            return result
    except RuntimeError as e:
        if "no running event loop" in str(e).lower():
            # No event loop, safe to use asyncio.run
            logger.info("[ASYNC_WRAPPER] No event loop detected, using asyncio.run")
            result = asyncio.run(async_func())
            logger.info("[ASYNC_WRAPPER] asyncio.run completed successfully (fallback)")
            return result
        else:
            logger.error(f"[ASYNC_WRAPPER] RuntimeError: {e}")
            raise e
    except Exception as e:
        logger.error(f"[ASYNC_WRAPPER] Unexpected error: {e}")
        raise e

# ================================
# AGENT TOOLS USING EXISTING SERVICES
# ================================

@tool
def text_summary_agent(query: str, file_data: Optional[Dict[str, Any]] = None) -> str:
    """
    Text summarization using DIRECT CALL to text_summary_node logic
    """
    try:
        logger.info(f"[TEXT_SUMMARY_AGENT] Processing: {query[:100]}...")
        
        # Import the actual node function
        from app.multi_agent.agents.conversation_agent.nodes.text_summary_node import _extract_text_from_message
        
        # Initialize services (giống node)
        text_service = TextSummaryService()
        
        # Extract text to summarize using node logic
        text_to_summarize = ""
        filename = "unknown"
        
        # Extract content from file if provided
        if file_data and file_data.get('raw_bytes'):
            try:
                raw_bytes = file_data.get('raw_bytes')
                content_type = file_data.get('content_type', '')
                filename = file_data.get('filename', 'unknown')
                
                logger.info(f"[TEXT_SUMMARY_AGENT] Processing file: {filename} ({content_type})")
                
                # Use existing extraction logic from helpers
                if content_type == "application/pdf":
                    from app.multi_agent.helpers.improved_pdf_extractor import ImprovedPDFExtractor
                    pdf_extractor = ImprovedPDFExtractor()
                    pdf_result = pdf_extractor.extract_text_from_pdf(raw_bytes)
                    text_to_summarize = pdf_result.get('text', '')
                    logger.info(f"[TEXT_SUMMARY_AGENT] Extracted PDF content: {len(text_to_summarize)} chars")
                    
                elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(raw_bytes))
                    text_to_summarize = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    logger.info(f"[TEXT_SUMMARY_AGENT] Extracted DOCX content: {len(text_to_summarize)} chars")
                    
                elif content_type.startswith("text/"):
                    text_to_summarize = raw_bytes.decode('utf-8')
                    logger.info(f"[TEXT_SUMMARY_AGENT] Extracted text content: {len(text_to_summarize)} chars")
                    
                else:
                    return f"❌ **Lỗi định dạng file**\n\nFile type {content_type} chưa được hỗ trợ."
                    
            except Exception as extract_error:
                logger.error(f"[TEXT_SUMMARY_AGENT] Content extraction error: {extract_error}")
                return f"❌ **Lỗi trích xuất nội dung**\n\nKhông thể đọc file {filename}: {str(extract_error)}"
        else:
            # Use EXACT node logic for text extraction
            text_to_summarize = _extract_text_from_message(query)
        
        # Check if we have text to summarize (EXACT node logic)
        if not text_to_summarize or len(text_to_summarize.strip()) < 10:
            return "Xin lỗi, tôi không tìm thấy văn bản nào để tóm tắt. Vui lòng cung cấp văn bản cần tóm tắt."
        
        # Use TextSummaryService với EXACT parameters từ node
        try:
            async def summarize_with_service():
                return await text_service.summarize_text(
                    text=text_to_summarize,
                    summary_type="general",
                    max_length=200,  # ← EXACT từ node (200)
                    language="vietnamese"
                )
            
            # Execute async function with safe wrapper
            summary_result = _run_async_safely(summarize_with_service)
            
            # Format response EXACT giống node
            response = f"📄 **Tóm tắt văn bản:**\n\n{summary_result['summary']}\n\n"
            response += f"📊 **Thống kê:** {summary_result['word_count']['original']} từ → {summary_result['word_count']['summary']} từ "
            response += f"(tỷ lệ nén: {summary_result['compression_ratio']})"
            
            if filename != "unknown":
                response = f"📄 **Tóm tắt tài liệu: {filename}**\n\n{summary_result['summary']}\n\n"
                response += f"📊 **Thống kê:** {summary_result['word_count']['original']} từ → {summary_result['word_count']['summary']} từ "
                response += f"(tỷ lệ nén: {summary_result['compression_ratio']})"
            
            logger.info("[TEXT_SUMMARY_AGENT] Successfully processed with DIRECT node logic")
            return response
            
        except Exception as e:
            logger.error(f"[TEXT_SUMMARY_AGENT] Summarization failed: {str(e)}")
            return f"Xin lỗi, có lỗi xảy ra khi tóm tắt văn bản: {str(e)}"
        
    except Exception as e:
        logger.error(f"[TEXT_SUMMARY_AGENT] Error: {str(e)}")
        return f"❌ **Lỗi xử lý tóm tắt**: {str(e)}"


@tool
def compliance_knowledge_agent(query: str, file_data: Optional[Dict[str, Any]] = None) -> str:
    """
    Compliance checking using DIRECT CALL to compliance/document endpoint
    Enhanced with better error handling and validation
    """
    try:
        logger.info(f"🔧 [COMPLIANCE_AGENT] TOOL CALLED with query: {query[:100]}...")
        
        # If file data is provided, use compliance/document endpoint DIRECTLY
        if file_data and file_data.get('raw_bytes'):
            logger.info(f"🔧 [COMPLIANCE_AGENT] Processing file: {file_data.get('filename')} ({len(file_data.get('raw_bytes', b''))} bytes)")
            
            # Validate file data
            if len(file_data.get('raw_bytes', b'')) == 0:
                return "❌ **Lỗi kiểm tra tuân thủ**: File rỗng hoặc không hợp lệ"
            
            # Import the EXACT service instead of endpoint
            from app.multi_agent.services.compliance_service import ComplianceValidationService
            from app.multi_agent.services.text_service import TextSummaryService
            
            try:
                # Initialize services
                compliance_service = ComplianceValidationService()
                text_service = TextSummaryService()
                
                # Extract text from document using text service
                raw_bytes = file_data.get('raw_bytes')
                filename = file_data.get('filename', 'document.pdf')
                content_type = file_data.get('content_type', 'application/pdf')
                
                # Get file extension
                import os
                file_extension = os.path.splitext(filename)[1].lower()
                file_size = len(raw_bytes)
                
                logger.info(f"🔧 [COMPLIANCE_AGENT] Processing file: {filename} ({file_size/1024:.1f}KB)")
                
                # Wrap both text extraction and compliance service in async function
                async def extract_and_validate():
                    # Extract text using text service (no timeout - same as original)
                    extracted_text = await text_service.extract_text_from_document(
                        file_content=raw_bytes,
                        file_extension=file_extension,
                        filename=filename
                    )
                    
                    if not extracted_text or len(extracted_text.strip()) < 50:
                        raise Exception("Không thể trích xuất đủ văn bản từ file để kiểm tra tuân thủ")
                    
                    logger.info(f"🔧 [COMPLIANCE_AGENT] Extracted {len(extracted_text)} characters from {filename}")
                    
                    # Call compliance service directly (no timeout - same as original)
                    logger.info(f"🔧 [COMPLIANCE_AGENT] Starting compliance validation")
                    
                    result = await compliance_service.validate_document_compliance(
                        ocr_text=extracted_text,
                        document_type=None  # Auto-detect
                    )
                    
                    logger.info(f"🔧 [COMPLIANCE_AGENT] Compliance validation completed")
                    return result
                
                # Execute async function with safe wrapper
                data = _run_async_safely(extract_and_validate)
                
                # Validate data exists
                if not data or not isinstance(data, dict):
                    logger.error(f"🔧 [COMPLIANCE_AGENT] Invalid or empty data: {type(data)}")
                    return "❌ **Lỗi kiểm tra tuân thủ**: Không thể xử lý tài liệu - dữ liệu không hợp lệ"
                
                logger.info(f"🔧 [COMPLIANCE_AGENT] Successfully got compliance data: {list(data.keys())}")
                
                # Return raw JSON data instead of formatted text
                import json
                try:
                    # Create response structure matching the endpoint format
                    response_data = {
                        "status": "success",
                        "data": data,
                        "message": f"Kiểm tra tuân thủ file {file_data.get('filename', 'unknown')} hoàn tất"
                    }
                    
                    # Return as formatted JSON string for better readability
                    json_response = json.dumps(response_data, ensure_ascii=False, indent=2)
                    
                    logger.info("🔧 [COMPLIANCE_AGENT] Successfully processed with DIRECT service call - returning JSON")
                    return json_response
                    
                except Exception as json_error:
                    logger.error(f"🔧 [COMPLIANCE_AGENT] JSON serialization error: {json_error}")
                    return f"❌ **Lỗi JSON serialization**: {str(json_error)}"
                
                # NOTE: Removed formatted response fallback - we only want JSON response
                    
            except Exception as service_error:
                logger.error(f"🔧 [COMPLIANCE_AGENT] Direct service call error: {service_error}")
                return f"❌ **Lỗi kiểm tra tuân thủ**: {str(service_error)}"
        
        else:
            # Handle text-based compliance queries using DIRECT node logic
            try:
                # Import the actual node functions
                from app.multi_agent.agents.conversation_agent.nodes.compliance_node import (
                    _determine_query_type,
                    _handle_regulation_query,
                    _handle_compliance_help,
                    _handle_general_compliance_chat,
                    _handle_compliance_help,
                    _handle_general_compliance_chat
                )
                
                # Use EXACT node logic for query type determination
                query_type = _determine_query_type(query)
                logger.info(f"🔧 [COMPLIANCE_AGENT] Query type determined: {query_type}")
                
                async def handle_compliance_query():
                    if query_type == "regulation_query":
                        return await _handle_regulation_query(query)
                    elif query_type == "compliance_help":
                        return await _handle_compliance_help(query)
                    else:
                        return await _handle_general_compliance_chat(query)
                
                # Execute async function with safe wrapper
                response = _run_async_safely(handle_compliance_query)
                
                logger.info("🔧 [COMPLIANCE_AGENT] Successfully processed with DIRECT node logic")
                return response
                
            except Exception as node_error:
                logger.error(f"🔧 [COMPLIANCE_AGENT] Node logic error: {node_error}")
                return f"❌ **Lỗi xử lý tuân thủ**: {str(node_error)}"

    except Exception as e:
        logger.error(f"🔧 [COMPLIANCE_AGENT] Tool error: {str(e)}")
        return f"❌ **Lỗi kiểm tra tuân thủ**: {str(e)}"

async def _handle_general_compliance_chat(message: str, compliance_service) -> str:
    """Handle general compliance chat (giống node logic)"""
    try:
        # Use general compliance validation
        result = await compliance_service.validate_compliance(
            document_content=message,
            compliance_standards=["UCP600", "ISBP821", "SBV"],
            validation_type="general"
        )
        
        if result and result.get('analysis'):
            return f"**Phân tích tuân thủ:**\n\n{result['analysis']}"
        else:
            return "Tôi có thể hỗ trợ bạn về các vấn đề tuân thủ ngân hàng. Vui lòng đặt câu hỏi cụ thể."
            
    except Exception as e:
        logger.error(f"Error in general compliance chat: {e}")
        return f"Có lỗi xảy ra: {str(e)}"

@tool
def risk_analysis_agent(query: str, file_data: Optional[Dict[str, Any]] = None) -> str:
    """
    Risk analysis using DIRECT CALL to existing risk API endpoint
    """
    try:
        logger.info(f"🔧 [RISK_AGENT] TOOL CALLED with query: {query[:100]}...")

        # Import required models and services
        from app.multi_agent.models.risk import RiskAssessmentRequest
        # Removed unused route imports - using service directly instead
        from fastapi import UploadFile
        import io

        # Extract basic info from query for risk assessment
        financial_data = _extract_basic_risk_data_from_query(query)
        
        # Extract text from file if provided
        if file_data and file_data.get('raw_bytes'):
            logger.info(f"🔧 [RISK_AGENT] Processing file: {file_data.get('filename')} ({len(file_data.get('raw_bytes', b''))} bytes)")
            
            try:
                # Extract text from file using the same logic as endpoint wrapper
                raw_bytes = file_data.get('raw_bytes')
                content_type = file_data.get('content_type', '')
                
                file_text = ""
                if content_type == "application/pdf":
                    from app.multi_agent.helpers.improved_pdf_extractor import ImprovedPDFExtractor
                    extractor = ImprovedPDFExtractor()
                    result = extractor.extract_text_from_pdf(raw_bytes)
                    file_text = result.get('text', '').strip()
                elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(raw_bytes))
                    file_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif content_type.startswith("text/"):
                    file_text = raw_bytes.decode('utf-8')
                
                financial_data['financial_documents'] = file_text
                logger.info(f"🔧 [RISK_AGENT] Extracted {len(file_text)} characters from file")
                
                if not file_text.strip():
                    logger.warning("🔧 [RISK_AGENT] No text extracted from file, proceeding with basic data")
                
            except Exception as file_error:
                logger.error(f"🔧 [RISK_AGENT] File processing error: {file_error}")
                return f"❌ **Lỗi xử lý file**: {str(file_error)}"
        
        # Call risk assessment with file content
        async def call_risk_api():
            from app.multi_agent.models.risk import RiskAssessmentRequest
            from app.multi_agent.services.risk_service import assess_risk
            
            risk_request = RiskAssessmentRequest(
                applicant_name=financial_data.get('applicant_name', 'Khách hàng'),
                business_type=financial_data.get('business_type', 'general'),
                requested_amount=financial_data.get('requested_amount', 1000000000),
                currency=financial_data.get('currency', 'VND'),
                loan_term=financial_data.get('loan_term', 12),
                loan_purpose=financial_data.get('loan_purpose', 'Kinh doanh'),
                assessment_type="comprehensive",
                collateral_type=financial_data.get('collateral_type', 'Không tài sản đảm bảo'),
                financial_documents=financial_data.get('financial_documents', '')  # ✅ Thêm file content
            )
            
            return await assess_risk(risk_request)
        
        # Execute async function with safe wrapper
        risk_result = _run_async_safely(call_risk_api)
        
        logger.info("🔧 [RISK_AGENT] Successfully processed with DIRECT service call")
        
        # Format response using EXACT existing API result
        if risk_result and hasattr(risk_result, 'data'):
            data = risk_result.data
            response = f"""📊 **Phân tích rủi ro - VPBank K-MULT**

**Thông tin đánh giá:**
• Tên khách hàng: {financial_data.get('applicant_name', 'Chưa xác định')}
• Số tiền yêu cầu: {financial_data.get('requested_amount', 0):,} VNĐ
• Loại hình kinh doanh: {financial_data.get('business_type', 'Chưa xác định')}

**Kết quả phân tích:**
• Điểm rủi ro: {data.get('risk_score', 'N/A')}
• Mức độ rủi ro: {data.get('risk_level', 'N/A')}
• Khuyến nghị: {data.get('recommendations', ['Cần đánh giá thêm'])[0] if data.get('recommendations') else 'Cần đánh giá thêm'}

**Báo cáo AI:**
{data.get('ai_report', 'Đang phân tích dữ liệu tài chính và đánh giá rủi ro...')}

---

*🤖 VPBank K-MULT Agent Studio*
*⏰ {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*"""
        else:
            # Fallback response
            response = f"""📊 **Phân tích rủi ro - VPBank K-MULT**

**Yêu cầu:** {query[:200]}...

**Phân tích sơ bộ:**
- Đang xử lý dữ liệu tài chính
- Áp dụng mô hình đánh giá rủi ro VPBank  
- Tuân thủ Basel III và quy định SBV

**Lưu ý:** Để có kết quả chính xác, vui lòng cung cấp:
• Tên khách hàng/doanh nghiệp
• Số tiền vay mong muốn
• Mục đích vay vốn
• Thông tin tài chính

---

*🤖 VPBank K-MULT Agent Studio*
*⏰ {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*"""
        
        return response
        
        # Format response using EXACT existing API result
        if risk_result and hasattr(risk_result, 'data'):
            data = risk_result.data
            response = f"""📊 **Phân tích rủi ro - VPBank K-MULT**

**Thông tin đánh giá:**
• Tên khách hàng: {financial_data.get('applicant_name', 'Chưa xác định')}
• Số tiền yêu cầu: {financial_data.get('requested_amount', 0):,} VNĐ
• Loại hình kinh doanh: {financial_data.get('business_type', 'Chưa xác định')}

**Kết quả phân tích:**
• Điểm rủi ro: {data.get('risk_score', 'N/A')}
• Mức độ rủi ro: {data.get('risk_level', 'N/A')}
• Khuyến nghị: {data.get('recommendations', ['Cần đánh giá thêm'])[0] if data.get('recommendations') else 'Cần đánh giá thêm'}

**Báo cáo AI:**
{data.get('ai_report', 'Đang phân tích dữ liệu tài chính và đánh giá rủi ro...')}

---

*🤖 VPBank K-MULT Agent Studio*
*⏰ {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*"""
        else:
            # Fallback response
            response = f"""📊 **Phân tích rủi ro - VPBank K-MULT**

**Yêu cầu:** {query[:200]}...

**Phân tích sơ bộ:**
- Đang xử lý dữ liệu tài chính
- Áp dụng mô hình đánh giá rủi ro VPBank  
- Tuân thủ Basel III và quy định SBV

**Lưu ý:** Để có kết quả chính xác, vui lòng cung cấp:
• Tên khách hàng/doanh nghiệp
• Số tiền vay mong muốn
• Mục đích vay vốn
• Thông tin tài chính

---

*🤖 VPBank K-MULT Agent Studio*
*⏰ {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*"""
        
        return response
        
    except Exception as e:
        logger.error(f"🔧 [RISK_AGENT] Tool error: {str(e)}")
        return f"❌ **Lỗi phân tích rủi ro**: {str(e)}"


def _extract_basic_risk_data_from_query(query: str) -> Dict[str, Any]:
    """Extract basic risk data from query - simplified version"""
    try:
        financial_data = {}
        query_lower = query.lower()
        
        # Extract applicant name (simple)
        if 'company' in query_lower or 'công ty' in query_lower:
            financial_data['applicant_name'] = 'Công ty ABC'
        else:
            financial_data['applicant_name'] = 'Khách hàng'
        
        # Extract amount (simple)
        import re
        amount_match = re.search(r'(\d+(?:,\d{3})*)', query)
        if amount_match:
            financial_data['requested_amount'] = int(amount_match.group(1).replace(',', ''))
        else:
            financial_data['requested_amount'] = 1000000000
        
        # Set defaults with proper structure for required fields
        financial_data.update({
            'business_type': 'general',
            'currency': 'VND',
            'loan_term': 12,
            'loan_purpose': 'Kinh doanh',
            'collateral_type': 'Không tài sản đảm bảo',
            # Required fields with proper structure
            'financials': {
                'revenue': 1000000000,
                'profit': 100000000,
                'assets': 2000000000,
                'liabilities': 500000000,
                'cash_flow': 300000000
            },
            'market_data': {
                'industry': 'general',
                'market_condition': 'stable',
                'competition_level': 'medium',
                'growth_potential': 'moderate'
            },
            'custom_factors': {
                'risk_tolerance': 'medium',
                'business_experience': 'established',
                'market_position': 'stable'
            }
        })
        
        return financial_data
        
    except Exception as e:
        logger.error(f"Error extracting basic risk data: {e}")
        return {
            'applicant_name': 'Khách hàng',
            'requested_amount': 1000000000,
            'business_type': 'general',
            'currency': 'VND',
            'loan_term': 12,
            'loan_purpose': 'Kinh doanh',
            'collateral_type': 'Không tài sản đảm bảo',
            # Required fields with proper structure
            'financials': {
                'revenue': 1000000000,
                'profit': 100000000,
                'assets': 2000000000,
                'liabilities': 500000000,
                'cash_flow': 300000000
            },
            'market_data': {
                'industry': 'general',
                'market_condition': 'stable',
                'competition_level': 'medium',
                'growth_potential': 'moderate'
            },
            'custom_factors': {
                'risk_tolerance': 'medium',
                'business_experience': 'established',
                'market_position': 'stable'
            }
        }


# ================================
# SUPERVISOR AGENT
# ================================

SUPERVISOR_PROMPT = """
You are a VPBank K-MULT INTELLIGENT routing supervisor. Your PRIMARY job is to route user requests to the correct specialized agent tool.

ROUTING RULES (STRICT PRIORITY ORDER):
1. COMPLIANCE QUERIES → compliance_knowledge_agent
   - Keywords: "kiểm tra", "tuân thủ", "compliance", "check", "validate", "verify", "quy định", "regulation", "UCP", "LC", "letter of credit"
   - File uploads for document validation
   - Banking regulation questions

2. TEXT SUMMARY QUERIES → text_summary_agent  
   - Keywords: "tóm tắt", "tóm lược", "tóm gọn", "summarize", "summary", "analyze document", "extract", "phân tích tài liệu"
   - File uploads for document summarization
   - Text analysis requests

3. RISK ANALYSIS QUERIES → risk_analysis_agent
   - Keywords: "phân tích", "rủi ro", "risk", "assess", "credit", "financial", "đánh giá", "tín dụng"
   - Financial analysis requests
   - Credit assessment queries

AGENT CAPABILITIES:
- compliance_knowledge_agent: Validates documents against UCP 600, ISBP 821, SBV regulations. Handles file uploads and compliance queries.
- text_summary_agent: Summarizes documents (PDF, DOCX, TXT) and text content. Provides statistics and compression ratios.
- risk_analysis_agent: Analyzes financial risk, credit assessment, Basel III compliance. Provides risk scores and recommendations.

ROUTING EXAMPLES:
User: "kiểm tra tuân thủ tài liệu LC" → compliance_knowledge_agent
User: "tóm tắt văn bản này" → text_summary_agent
User: "phân tích rủi ro tín dụng" → risk_analysis_agent
User: "UCP 600 quy định gì?" → compliance_knowledge_agent
User: "đánh giá khả năng trả nợ" → risk_analysis_agent

CRITICAL INSTRUCTIONS:
- You MUST call exactly ONE tool for every request
- You CANNOT provide direct answers or explanations
- You MUST analyze the user's intent and route to the appropriate tool
- If unclear, default to compliance_knowledge_agent for banking-related queries

Your response should ONLY be the tool execution result. No additional commentary.

EMERGENCY PROTOCOL:
If you cannot determine which tool to use, call compliance_knowledge_agent as default.

YOUR RESPONSE MUST BE: Tool execution result ONLY. No preamble, no explanation, no apology.
"""

# Create supervisor with stronger model configuration
supervisor_agent = Agent(
    system_prompt=SUPERVISOR_PROMPT,
    tools=[text_summary_agent, compliance_knowledge_agent, risk_analysis_agent],
    model=BedrockModel(
        model_id=BEDROCK_MODEL_ID,
        boto_session=boto_session,
        temperature=0.1,  # Lower temperature for more deterministic behavior
        top_p=0.8,
        streaming=False,  # Disable streaming for more reliable tool calls
        max_tokens=1000   # Limit tokens to force concise responses
    )
)

# ================================
# MAIN SYSTEM CLASS
# ================================

class PureStrandsVPBankSystem:
    """VPBank K-MULT Agent Studio - Clean Pure Strands Implementation with DIRECT NODE INTEGRATION"""
    
    def __init__(self):
        self.supervisor = supervisor_agent
        self.session_data = {}
        self.processing_stats = {
            "total_requests": 0,
            "successful_responses": 0,
            "errors": 0,
            "agent_usage": {
                "text_summary_agent": 0,
                "compliance_knowledge_agent": 0,
                "risk_analysis_agent": 0,
                "general_redirect": 0
            }
        }
    
    def _is_banking_related(self, query: str) -> bool:
        """
        Smart banking relevance detection with pre-filtering
        Returns True if query is banking/finance related, False otherwise
        """
        try:
            query_lower = query.lower().strip()
            
            # Empty or very short queries - allow through
            if len(query_lower) < 3:
                return True
            
            # Strong non-banking indicators (high confidence)
            non_banking_strong = [
                # Weather & Environment
                'thời tiết', 'weather', 'nhiệt độ', 'temperature', 'mưa', 'rain', 'nắng', 'sunny',
                
                # Food & Cooking
                'nấu ăn', 'cooking', 'recipe', 'công thức', 'món ăn', 'food', 'ăn uống',
                'nhà hàng', 'restaurant', 'quán ăn', 'đồ ăn',
                
                # Travel & Tourism
                'du lịch', 'travel', 'tour', 'khách sạn', 'hotel', 'máy bay', 'flight',
                'vé máy bay', 'booking', 'đặt phòng', 'resort',
                
                # Sports & Entertainment
                'thể thao', 'sports', 'bóng đá', 'football', 'tennis', 'basketball',
                'phim', 'movie', 'cinema', 'âm nhạc', 'music', 'ca sĩ', 'singer',
                'game', 'gaming', 'chơi game', 'video game',
                
                # Health & Medical
                'sức khỏe', 'health', 'y tế', 'medical', 'bác sĩ', 'doctor', 'bệnh viện', 'hospital',
                'thuốc', 'medicine', 'điều trị', 'treatment',
                
                # Technology (non-fintech)
                'điện thoại', 'phone', 'smartphone', 'laptop', 'computer', 'máy tính',
                'internet', 'wifi', 'facebook', 'instagram', 'tiktok',
                
                # Education (non-finance)
                'học tập', 'study', 'trường học', 'school', 'đại học', 'university',
                'bài tập', 'homework', 'thi cử', 'exam',
                
                # Personal & Lifestyle
                'tình yêu', 'love', 'hẹn hò', 'dating', 'gia đình', 'family',
                'mua sắm', 'shopping', 'thời trang', 'fashion', 'làm đẹp', 'beauty',
                
                # Stock Market (non-banking specific)
                'giá cả cổ phiếu', 'tình hình cổ phiếu', 'thị trường chứng khoán hôm nay',
                'cổ phiếu tăng giảm', 'biến động thị trường', 'giá cổ phiếu hôm nay',
                
                # Commodity Prices (non-banking)
                'giá vàng hôm nay', 'giá vàng', 'tình hình giá vàng', 'vàng tăng giá',
                'giá dầu', 'giá dầu hôm nay', 'giá xăng', 'giá USD', 'tỷ giá hôm nay',
                'giá bitcoin', 'giá crypto', 'tiền điện tử'
            ]
            
            # Check for strong non-banking indicators
            for keyword in non_banking_strong:
                if keyword in query_lower:
                    logger.info(f"[PRE_FILTER] Non-banking keyword detected: '{keyword}' in query")
                    return False
            
            # Banking/Finance keywords (comprehensive but specific to banking services)
            banking_keywords = [
                # Core Banking Services
                'ngân hàng', 'bank', 'banking', 'vpbank', 'vp bank',
                'tài khoản', 'account', 'số dư', 'balance', 'giao dịch', 'transaction',
                'chuyển khoản', 'transfer', 'rút tiền', 'withdraw', 'gửi tiền', 'deposit',
                
                # Credit & Loans (Banking specific)
                'tín dụng', 'credit', 'vay', 'loan', 'cho vay', 'lending',
                'lãi suất', 'interest rate', 'thế chấp', 'mortgage', 'bảo lãnh', 'guarantee',
                'khoản vay', 'loan amount', 'trả nợ', 'repayment',
                
                # Banking Finance (not stock market)
                'tài chính ngân hàng', 'banking finance', 'dịch vụ tài chính', 'financial services',
                'sản phẩm ngân hàng', 'banking products', 'tiền gửi', 'savings',
                
                # Investment Banking (not stock trading)
                'ngân hàng đầu tư', 'investment banking', 'tư vấn tài chính', 'financial advisory',
                'quản lý tài sản', 'asset management',
                
                # Risk & Compliance (Banking specific)
                'rủi ro tín dụng', 'credit risk', 'đánh giá rủi ro', 'risk assessment',
                'tuân thủ', 'compliance', 'quy định ngân hàng', 'banking regulation',
                'kiểm tra', 'check', 'validate', 'verify', 'xác minh',
                
                # Trade Finance (Banking specific)
                'lc', 'letter of credit', 'thư tín dụng', 'ucp', 'ucp 600',
                'isbp', 'bill of lading', 'vận đơn', 'xuất nhập khẩu', 'export', 'import',
                'tài chính thương mại', 'trade finance',
                
                # Document Processing (Banking context)
                'tóm tắt tài liệu', 'document summary', 'phân tích báo cáo', 'report analysis',
                'tài liệu ngân hàng', 'banking document', 'báo cáo tài chính', 'financial report',
                'trích xuất', 'extract', 'xử lý tài liệu', 'document processing',
                
                # Regulatory Bodies
                'sbv', 'nhnn', 'basel', 'basel iii', 'central bank', 'ngân hàng trung ương',
                'quy định sbv', 'sbv regulation',
                
                # Business Banking
                'doanh nghiệp', 'enterprise', 'công ty', 'company', 'business banking',
                'tài chính doanh nghiệp', 'corporate finance', 'thương mại', 'commercial banking'
            ]
            
            # Check for banking keywords
            for keyword in banking_keywords:
                if keyword in query_lower:
                    logger.info(f"[PRE_FILTER] Banking keyword detected: '{keyword}' in query")
                    return True
            
            # Check for file upload context (usually banking documents)
            file_extensions = ['.pdf', '.docx', '.doc', '.txt', '.xlsx']
            if any(ext in query_lower for ext in file_extensions):
                logger.info("[PRE_FILTER] File extension detected - assuming banking document")
                return True
            
            # Ambiguous cases - allow through (better false positive than negative)
            # This ensures we don't accidentally block legitimate banking questions
            logger.info(f"[PRE_FILTER] Ambiguous query - allowing through: '{query_lower[:50]}...'")
            return True
            
        except Exception as e:
            logger.error(f"[PRE_FILTER] Error in banking detection: {e}")
            # On error, allow through to be safe
            return True
    
    def _get_redirect_message(self, query: str = "") -> str:
        """
        Generate interactive redirect message based on query context
        """
        query_lower = query.lower().strip()
        
        # Detect topic and create contextual response
        if any(keyword in query_lower for keyword in ['thời tiết', 'weather', 'mưa', 'nắng', 'nhiệt độ']):
            topic_response = "Rất tiếc, tôi chuyên xử lý các vấn đề ngân hàng nên thông tin thời tiết nằm ngoài hiểu biết của tôi."
            
        elif any(keyword in query_lower for keyword in ['giá vàng', 'giá dầu', 'giá cổ phiếu', 'bitcoin', 'crypto']):
            topic_response = "Tôi hiểu bạn quan tâm đến thông tin thị trường, nhưng tôi chuyên về dịch vụ ngân hàng nên không thể cung cấp giá cả hàng hóa hay chứng khoán."
            
        elif any(keyword in query_lower for keyword in ['nấu ăn', 'món ăn', 'recipe', 'cooking', 'nhà hàng']):
            topic_response = "Tôi thấy bạn hỏi về ẩm thực! Tuy nhiên, tôi là trợ lý chuyên về ngân hàng nên không thể tư vấn về nấu ăn."
            
        elif any(keyword in query_lower for keyword in ['du lịch', 'travel', 'khách sạn', 'tour', 'máy bay']):
            topic_response = "Du lịch thật thú vị! Nhưng tôi chuyên hỗ trợ các dịch vụ ngân hàng nên không thể tư vấn về du lịch."
            
        elif any(keyword in query_lower for keyword in ['phim', 'movie', 'âm nhạc', 'music', 'game']):
            topic_response = "Tôi hiểu bạn quan tâm đến giải trí, nhưng chuyên môn của tôi là về ngân hàng và tài chính."
            
        elif any(keyword in query_lower for keyword in ['sức khỏe', 'health', 'bác sĩ', 'bệnh viện', 'thuốc']):
            topic_response = "Sức khỏe rất quan trọng! Tuy nhiên, tôi chuyên về lĩnh vực ngân hàng nên không thể tư vấn y tế."
            
        elif any(keyword in query_lower for keyword in ['học tập', 'study', 'trường học', 'bài tập', 'thi cử']):
            topic_response = "Học tập là điều tuyệt vời! Nhưng tôi chuyên hỗ trợ các vấn đề ngân hàng nên không thể giúp về học tập."
            
        elif any(keyword in query_lower for keyword in ['tình yêu', 'love', 'hẹn hò', 'dating', 'gia đình']):
            topic_response = "Tôi hiểu những vấn đề cá nhân rất quan trọng, nhưng tôi chuyên về dịch vụ ngân hàng."
            
        else:
            # Generic response for unrecognized topics
            topic_response = f"Tôi thấy bạn hỏi về '{query[:50]}...'. Tuy nhiên, tôi chuyên hỗ trợ các vấn đề ngân hàng và tài chính."
        
        return f"""💬 **{topic_response}**

🏦 **Tôi có thể giúp bạn với:**
• 📄 **Tóm tắt tài liệu** - Phân tích báo cáo, hợp đồng, văn bản
• ⚖️ **Kiểm tra tuân thủ** - UCP 600, quy định SBV, ISBP 821  
• 📊 **Phân tích rủi ro** - Đánh giá tín dụng, Basel III
• 💳 **Letter of Credit** - Xử lý thư tín dụng, tài liệu thương mại

💡 **Thử hỏi tôi:**
- "Tóm tắt báo cáo này"
- "Kiểm tra tuân thủ tài liệu LC"  
- "Phân tích rủi ro khoản vay 10 tỷ"
- "UCP 600 quy định gì về vận đơn?"

Bạn có câu hỏi nào về ngân hàng không? 😊"""
    
    async def process_request(
        self, 
        user_message: str, 
        conversation_id: str,
        context: Optional[Dict[str, Any]] = None,
        uploaded_file: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process user request with PRE-FILTERING + MANUAL ROUTING + DIRECT NODE INTEGRATION"""
        try:
            self.processing_stats["total_requests"] += 1
            start_time = datetime.now()
            
            logger.info(f"[PURE_STRANDS] Processing request for conversation {conversation_id}")
            logger.info(f"[DEBUG] PRE-FILTERING: About to check banking relevance for: '{user_message[:50]}...'")
            
            # ================================
            # PRE-FILTERING: Check if banking-related
            # ================================
            
            # Skip pre-filtering if file is uploaded (assume banking document)
            if not uploaded_file and not self._is_banking_related(user_message):
                logger.info(f"[PRE_FILTER] Non-banking query detected: '{user_message[:100]}...'")
                
                processing_time = (datetime.now() - start_time).total_seconds()
                self.processing_stats["successful_responses"] += 1
                self.processing_stats["agent_usage"]["general_redirect"] += 1
                
                # Store session data
                self.session_data[conversation_id] = {
                    "last_message": user_message,
                    "last_response": "general_redirect",
                    "agent_used": "general_redirect",
                    "timestamp": datetime.now().isoformat(),
                    "processing_time": processing_time,
                    "file_processed": None
                }
                
                return {
                    "status": "success",
                    "conversation_id": conversation_id,
                    "response": self._get_redirect_message(user_message),
                    "agent_used": "general_redirect",
                    "processing_time": processing_time,
                    "timestamp": datetime.now().isoformat(),
                    "system": "pure_strands_vpbank_pre_filter",
                    "file_processed": None,
                    "request_type": "non_banking_redirect"
                }
            
            logger.info("[PRE_FILTER] Banking-related query confirmed - proceeding with agent routing")
            
            # ================================
            # ENHANCED MANUAL ROUTING - Primary approach for reliability with DIRECT NODE CALLS
            # ================================
            message_lower = user_message.lower()
            selected_agent = None
            
            # Enhanced keyword detection with priority scoring
            compliance_keywords = [
                'kiểm tra', 'tuân thủ', 'compliance', 'check', 'validate', 'verify', 'conform',
                'quy định', 'regulation', 'ucp', 'isbp', 'sbv', 'letter of credit', 'lc',
                'banking regulation', 'document validation', 'compliance check'
            ]
            
            summary_keywords = [
                'tóm tắt', 'summarize', 'summary', 'analyze document', 'extract', 'document analysis',
                'phân tích tài liệu', 'trích xuất', 'tổng hợp', 'rút gọn', 'document summary'
            ]
            
            risk_keywords = [
                'phân tích rủi ro', 'rủi ro', 'risk', 'analysis', 'credit', 'assess', 'financial',
                'đánh giá', 'tín dụng', 'credit assessment', 'risk analysis', 'financial analysis',
                'basel', 'credit score', 'loan assessment'
            ]
            
            # Calculate keyword match scores
            compliance_score = sum(1 for keyword in compliance_keywords if keyword in message_lower)
            summary_score = sum(1 for keyword in summary_keywords if keyword in message_lower)
            risk_score = sum(1 for keyword in risk_keywords if keyword in message_lower)
            
            # Determine primary intent based on highest score
            max_score = max(compliance_score, summary_score, risk_score)
            
            if max_score > 0:
                if compliance_score == max_score:
                    selected_agent = "compliance"
                    logger.info(f"[PURE_STRANDS] Manual routing: COMPLIANCE detected (score: {compliance_score})")
                elif summary_score == max_score:
                    selected_agent = "summary"
                    logger.info(f"[PURE_STRANDS] Manual routing: SUMMARY detected (score: {summary_score})")
                elif risk_score == max_score:
                    selected_agent = "risk"
                    logger.info(f"[PURE_STRANDS] Manual routing: RISK detected (score: {risk_score})")
            
            # Special handling for file uploads
            if uploaded_file and not selected_agent:
                file_ext = uploaded_file.get('filename', '').lower().split('.')[-1]
                if file_ext in ['pdf', 'docx', 'txt']:
                    # Default to compliance for banking documents
                    selected_agent = "compliance"
                    logger.info("[PURE_STRANDS] Manual routing: FILE UPLOAD → defaulting to COMPLIANCE")
            
            # Execute single agent with MANUAL ROUTING + DIRECT NODE CALLS (Primary approach)
            if selected_agent:
                logger.info(f"[PURE_STRANDS] Using MANUAL routing to {selected_agent} agent with DIRECT node integration")
                
                try:
                    if selected_agent == "compliance":
                        response = compliance_knowledge_agent(user_message, file_data=uploaded_file)
                        agent_used = "compliance_knowledge_agent"
                    elif selected_agent == "summary":
                        response = text_summary_agent(user_message, file_data=uploaded_file)
                        agent_used = "text_summary_agent"
                    elif selected_agent == "risk":
                        response = risk_analysis_agent(user_message, file_data=uploaded_file)
                        agent_used = "risk_analysis_agent"
                    
                    logger.info(f"[PURE_STRANDS] Manual routing successful with DIRECT node integration: {agent_used}")
                    
                    # Validate response is not empty
                    if not response or len(str(response).strip()) < 10:
                        logger.error(f"[PURE_STRANDS] Empty response from {agent_used}, falling back to Strands")
                        raise Exception("Empty response from manual routing")
                        
                except Exception as manual_error:
                    logger.error(f"[PURE_STRANDS] Manual routing failed: {manual_error}")
                    # Fallback to Strands if manual routing fails
                    selected_agent = None
            
            # Fallback to Strands supervisor ONLY if manual routing failed or unclear intent
            if not selected_agent:
                logger.info("[PURE_STRANDS] Using Strands supervisor as fallback")
                
                try:
                    if uploaded_file:
                        logger.info(f"[PURE_STRANDS] Creating file-aware tools for: {uploaded_file.get('filename')}")
                        
                        # Create bound tools with file data
                        @tool
                        def text_summary_with_file(query: str) -> str:
                            return text_summary_agent(query, file_data=uploaded_file)
                        
                        @tool  
                        def compliance_with_file(query: str) -> str:
                            return compliance_knowledge_agent(query, file_data=uploaded_file)
                        
                        @tool
                        def risk_analysis_with_file(query: str) -> str:
                            return risk_analysis_agent(query, file_data=uploaded_file)
                        
                        # Create file-aware supervisor
                        file_supervisor = Agent(
                            system_prompt=SUPERVISOR_PROMPT,
                            tools=[text_summary_with_file, compliance_with_file, risk_analysis_with_file],
                            model=BedrockModel(
                                model_id=BEDROCK_MODEL_ID,
                                boto_session=boto_session,
                                temperature=0.1,
                                top_p=0.8,
                                streaming=False,
                                max_tokens=1000
                            )
                        )
                        
                        response = file_supervisor(user_message)
                        logger.info("[PURE_STRANDS] Used file-aware supervisor")
                        
                    else:
                        # Use regular supervisor
                        response = self.supervisor(user_message)
                        logger.info("[PURE_STRANDS] Used regular supervisor")
                    
                    agent_used = self._detect_agent_used(str(response))
                    
                    # Validate Strands response
                    if not response or len(str(response).strip()) < 10:
                        logger.error("[PURE_STRANDS] Empty response from Strands supervisor")
                        response = "❌ **Lỗi hệ thống**: Không thể xử lý yêu cầu. Vui lòng thử lại."
                        agent_used = "error_fallback"
                        
                except Exception as strands_error:
                    logger.error(f"[PURE_STRANDS] Strands supervisor failed: {strands_error}")
                    response = "❌ **Lỗi hệ thống**: Không thể xử lý yêu cầu. Vui lòng thử lại."
                    agent_used = "error_fallback"
            
            processing_time = (datetime.now() - start_time).total_seconds()
            self.processing_stats["successful_responses"] += 1
            
            # Update agent usage stats
            if agent_used in self.processing_stats["agent_usage"]:
                self.processing_stats["agent_usage"][agent_used] += 1
            
            # Store session data
            self.session_data[conversation_id] = {
                "last_message": user_message,
                "last_response": str(response),
                "agent_used": agent_used,
                "timestamp": datetime.now().isoformat(),
                "processing_time": processing_time,
                "file_processed": uploaded_file.get('filename') if uploaded_file else None
            }
            
            result = {
                "status": "success",
                "conversation_id": conversation_id,
                "response": str(response),
                "agent_used": agent_used,
                "processing_time": processing_time,
                "timestamp": datetime.now().isoformat(),
                "system": "pure_strands_vpbank_manual_routing",
                "file_processed": uploaded_file.get('filename') if uploaded_file else None
            }
            
            logger.info(f"[PURE_STRANDS] Successfully processed in {processing_time:.2f}s using {agent_used}")
            return result
            
        except Exception as e:
            self.processing_stats["errors"] += 1
            logger.error(f"[PURE_STRANDS] Error processing request: {str(e)}")
            
            return {
                "status": "error",
                "conversation_id": conversation_id,
                "response": f"Error processing request: {str(e)}",
                "agent_used": "error_handler",
                "processing_time": 0,
                "timestamp": datetime.now().isoformat(),
                "system": "pure_strands_vpbank_manual_routing",
                "error": str(e)
            }
    
    def _detect_agent_used(self, response: str) -> str:
        """Detect which agent was used based on response content and logging"""
        response_lower = response.lower()
        
        # Check for agent-specific markers in response
        if "⚖️" in response or "kiểm tra tuân thủ" in response_lower or "compliance" in response_lower:
            return "compliance_knowledge_agent"
        elif "📄" in response or "tóm tắt" in response_lower or "summary" in response_lower:
            return "text_summary_agent"
        elif "📊" in response or "phân tích rủi ro" in response_lower or "risk" in response_lower:
            return "risk_analysis_agent"
        else:
            # Check for generic supervisor responses (failed tool execution)
            generic_phrases = [
                "i apologize", "i'll help", "let me", "would you like",
                "there was an issue", "could be due to", "try again"
            ]
            
            if any(phrase in response_lower for phrase in generic_phrases):
                logger.warning(f"[SUPERVISOR] Detected generic response instead of tool execution: {response[:100]}...")
                return "supervisor_direct_failed"
            elif any(marker in response for marker in ["**", "•", "---", "VPBank"]):
                # Likely from a tool but couldn't identify which one
                return "unknown_tool"
            else:
                return "supervisor_direct"
    
    def get_system_status(self) -> Dict[str, Any]:
        """Get system status with PRE-FILTERING + DIRECT NODE INTEGRATION info"""
        return {
            "system": "VPBank K-MULT Pure Strands with PRE-FILTERING + DIRECT NODE INTEGRATION",
            "supervisor_status": "active",
            "pre_filtering": {
                "enabled": True,
                "description": "Smart banking relevance detection",
                "non_banking_handling": "Friendly redirect with capability overview",
                "banking_keywords": "Comprehensive banking/finance vocabulary",
                "redirect_count": self.processing_stats["agent_usage"].get("general_redirect", 0)
            },
            "available_agents": [
                "text_summary_agent (→ text_summary_node DIRECT)",
                "compliance_knowledge_agent (→ compliance_node DIRECT)", 
                "risk_analysis_agent (→ risk API DIRECT)",
                "general_redirect (→ pre-filter redirect)"
            ],
            "node_integration": {
                "text_summary_agent": "Uses text_summary_node._extract_text_from_message + TextSummaryService",
                "compliance_knowledge_agent": "Uses compliance_node functions (_determine_query_type, _handle_regulation_query, etc.)",
                "risk_analysis_agent": "Uses risk_routes.assess_risk_endpoint DIRECTLY",
                "general_redirect": "Pre-filtering with smart banking relevance detection"
            },
            "routing_flow": [
                "1. Pre-filtering: Banking relevance check",
                "2. Manual routing: Keyword-based agent selection", 
                "3. Strands supervisor: AI-powered fallback",
                "4. Direct node integration: Service calls"
            ],
            "active_sessions": len(self.session_data),
            "processing_stats": self.processing_stats,
            "last_updated": datetime.now().isoformat()
        }

# ================================
# GLOBAL INSTANCE - FORCE RECREATION
# ================================

# Force recreation of the system instance to ensure latest code is used
pure_strands_vpbank_system = PureStrandsVPBankSystem()

async def process_pure_strands_request(user_message: str, conversation_id: str, context: Optional[Dict] = None, uploaded_file: Optional[Dict] = None):
    """
    Process request through Pure Strands system with PRE-FILTERING
    This function ensures the latest instance with pre-filtering is used
    """
    logger.info(f"[WRAPPER] Processing request: '{user_message[:50]}...'")
    return await pure_strands_vpbank_system.process_request(user_message, conversation_id, context, uploaded_file)

def get_pure_strands_system_status():
    return pure_strands_vpbank_system.get_system_status()

__all__ = [
    "pure_strands_vpbank_system",
    "process_pure_strands_request", 
    "get_pure_strands_system_status"
]
